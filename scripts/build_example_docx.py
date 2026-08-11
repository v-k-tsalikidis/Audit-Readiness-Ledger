"""Regenerate the Word document in the example set.

The example set needs one real .docx so that every run exercises the Word
reading path and the paragraph-based citations, not only the unit tests. A
Word file is a zip archive and does not belong in a diff, so it is built
from this script rather than edited by hand.

    python scripts/build_example_docx.py

The control block is a table because that is where a policy really puts it,
and reading tables is the part of the Word path most likely to break.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

TARGET = Path("examples/meltemi-logistics/documents/logging-and-monitoring-sop.docx")

CONTROL_BLOCK = (
    ("Document owner", "Nikos Alevras, Systems Administrator"),
    ("Approved by", "Head of IT"),
    ("Version", "1.3"),
    ("Approval date", "2026-01-20"),
    ("Classification", "Internal"),
    ("Next review", "2027-01-20"),
)

BODY = (
    "Logging",
    (
        "Audit logs are generated on the transport management system, the customer "
        "portal and the firewall. Event logs are retained for one year and are "
        "protected from alteration by staff who administer the systems they cover."
    ),
    "Monitoring activities",
    (
        "The firewall and server activity is monitored for anomalous behaviour. "
        "Alerts are reviewed each working morning by the systems administrator."
    ),
    "Clock synchronisation",
    (
        "All servers take their time source from NTP so that logs from different "
        "systems can be compared accurately."
    ),
    "Network security",
    (
        "Firewall rules are documented and reviewed twice a year. Filtering is "
        "applied at the network perimeter."
    ),
    "Related documents",
    # Deliberately dangling: no Records Retention Standard is supplied.
    "Retention periods are defined in accordance with the Records Retention Standard.",
)


def build(target: Path = TARGET) -> Path:
    document = Document()
    document.add_heading("Logging and Monitoring Standard Operating Procedure", 0)

    table = document.add_table(rows=len(CONTROL_BLOCK), cols=2)
    for row, (label, value) in enumerate(CONTROL_BLOCK):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = value

    for paragraph in BODY:
        document.add_paragraph(paragraph)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    return target


if __name__ == "__main__":
    print(f"wrote {build()}")
