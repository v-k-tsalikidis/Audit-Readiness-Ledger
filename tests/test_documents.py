"""Tests for reading a folder of documents.

The cases that matter are the ones where being wrong is quiet: a citation
that points at the wrong place, a term picked up from a code sample, and a
file that failed to open without anyone being told.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from audit_readiness_ledger.documents import (
    MAX_BYTES,
    load_directory,
    load_document,
)


class Folder:
    """A throwaway folder, so tests touch the real filesystem."""

    def __init__(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.path = Path(self._tmp.name)

    def write(self, name: str, text: str, encoding: str = "utf-8") -> Path:
        target = self.path / name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(text.encode(encoding))
        return target

    def write_bytes(self, name: str, data: bytes) -> Path:
        target = self.path / name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(data)
        return target

    def close(self):
        self._tmp.cleanup()


class CitationsLeadSomewhere(unittest.TestCase):
    def setUp(self):
        self.folder = Folder()
        self.addCleanup(self.folder.close)

    def test_line_numbers_survive_noise_removal(self):
        """Blanked regions keep their slots so later lines keep their numbers."""
        self.folder.write(
            "p.md",
            "# Title\n```\nbackup --force\n```\nBackups are tested quarterly.\n",
        )
        document, unit = load_document(self.folder.path / "p.md")
        self.assertEqual(unit, "line")
        self.assertEqual(len(document.lines), 5)
        # The real sentence is still on line 5, where a reader will find it.
        self.assertIn("tested quarterly", document.lines[4])

    def test_a_term_inside_a_code_fence_is_not_read(self):
        self.folder.write("p.md", "# T\n```\nbackup restore\n```\n")
        document, _ = load_document(self.folder.path / "p.md")
        self.assertNotIn("backup", document.text.lower())

    def test_front_matter_is_ignored_but_only_at_the_top(self):
        self.folder.write("p.md", "---\nowner: nobody\n---\nBackups are tested.\n")
        document, _ = load_document(self.folder.path / "p.md")
        self.assertNotIn("nobody", document.text)
        self.assertIn("Backups are tested.", document.text)

    def test_a_horizontal_rule_mid_document_is_not_front_matter(self):
        self.folder.write("p.md", "Intro.\n---\nBackups are tested.\n")
        document, _ = load_document(self.folder.path / "p.md")
        self.assertIn("Backups are tested.", document.text)


class WordDocumentsAreCitedAsParagraphs(unittest.TestCase):
    def setUp(self):
        self.folder = Folder()
        self.addCleanup(self.folder.close)

    def _docx(self, name: str, paragraphs: list[str], table: list[list[str]] | None = None):
        from docx import Document as DocxDocument

        document = DocxDocument()
        for text in paragraphs:
            document.add_paragraph(text)
        if table:
            added = document.add_table(rows=len(table), cols=len(table[0]))
            for row_index, row in enumerate(table):
                for cell_index, value in enumerate(row):
                    added.cell(row_index, cell_index).text = value
        path = self.folder.path / name
        document.save(str(path))
        return path

    def test_the_unit_is_paragraph_not_line(self):
        path = self._docx("p.docx", ["Backups run nightly.", "Recovery is tested."])
        _, unit = load_document(path)
        self.assertEqual(unit, "paragraph")

    def test_table_cells_are_read(self):
        """The document control block is a table, and the hygiene checks need it."""
        path = self._docx(
            "p.docx",
            ["Backup Policy"],
            table=[["Owner", "Jane Doe"], ["Next review", "2027-01-01"]],
        )
        document, _ = load_document(path)
        self.assertIn("Jane Doe", document.text)
        self.assertIn("2027-01-01", document.text)

    def test_a_row_stays_on_one_line(self):
        path = self._docx("p.docx", ["Policy"], table=[["Owner", "Jane Doe"]])
        document, _ = load_document(path)
        self.assertTrue(any(line == "Owner | Jane Doe" for line in document.lines))


class NothingIsSkippedSilently(unittest.TestCase):
    def setUp(self):
        self.folder = Folder()
        self.addCleanup(self.folder.close)

    def test_an_unreadable_format_is_reported(self):
        self.folder.write("scan.pdf", "not really a pdf")
        result = load_directory(self.folder.path)
        self.assertEqual(len(result.documents), 0)
        self.assertEqual(len(result.skipped), 1)
        self.assertIn(".pdf", result.skipped[0].reason)

    def test_a_corrupt_docx_is_reported_not_crashed_on(self):
        self.folder.write_bytes("broken.docx", b"PK\x03\x04 truncated")
        result = load_directory(self.folder.path)
        self.assertEqual(len(result.skipped), 1)
        self.assertIn("docx", result.skipped[0].reason.lower())

    def test_an_empty_file_is_reported(self):
        self.folder.write("empty.md", "")
        result = load_directory(self.folder.path)
        self.assertIn("empty", result.skipped[0].reason)

    def test_a_file_of_only_whitespace_is_reported(self):
        self.folder.write("blank.md", "\n\n   \n")
        result = load_directory(self.folder.path)
        self.assertEqual(len(result.documents), 0)
        self.assertIn("no readable text", result.skipped[0].reason)

    def test_an_oversized_file_is_reported(self):
        self.folder.write("huge.txt", "x" * (MAX_BYTES + 1))
        result = load_directory(self.folder.path)
        self.assertIn("larger than", result.skipped[0].reason)

    def test_word_lock_files_and_hidden_files_are_passed_over_quietly(self):
        """These are noise, not findings, so they do not clutter the report."""
        self.folder.write("~$policy.docx", "lock")
        self.folder.write(".hidden/notes.md", "Backups.")
        result = load_directory(self.folder.path)
        self.assertEqual(result.documents, ())
        self.assertEqual(result.skipped, ())


class ReadingIsRepeatable(unittest.TestCase):
    def setUp(self):
        self.folder = Folder()
        self.addCleanup(self.folder.close)

    def test_documents_come_back_in_a_stable_order(self):
        for name in ["c.md", "a.md", "b.md"]:
            self.folder.write(name, "Backups are tested.")
        first = [d.name for d in load_directory(self.folder.path).documents]
        second = [d.name for d in load_directory(self.folder.path).documents]
        self.assertEqual(first, second)
        self.assertEqual(first, ["a.md", "b.md", "c.md"])

    def test_nested_folders_keep_their_path_in_the_name(self):
        self.folder.write("policies/access.md", "Access control is defined.")
        result = load_directory(self.folder.path)
        self.assertEqual(result.documents[0].name, "policies/access.md")

    def test_a_greek_windows_encoding_is_read(self):
        """Real Greek document sets are not all UTF-8."""
        self.folder.write("greek.txt", "Πολιτική ασφάλειας. Backups are tested.", "cp1253")
        result = load_directory(self.folder.path)
        self.assertEqual(len(result.documents), 1)
        self.assertIn("Backups are tested.", result.documents[0].text)


class MissingInputFailsClearly(unittest.TestCase):
    def test_a_folder_that_does_not_exist(self):
        with self.assertRaises(FileNotFoundError):
            load_directory(Path("/nonexistent/folder/xyz"))

    def test_a_file_given_where_a_folder_was_expected(self):
        folder = Folder()
        self.addCleanup(folder.close)
        path = folder.write("p.md", "text")
        with self.assertRaises(NotADirectoryError):
            load_directory(path)


if __name__ == "__main__":
    unittest.main()
