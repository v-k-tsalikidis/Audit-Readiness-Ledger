"""Read a folder of policy documents into text the matcher can cite.

Three decisions here decide whether the output is useful or misleading.

**The citation has to lead somewhere.** A finding says "backup-policy.md
line 14", and the reader opens the file at line 14 and sees it. Markdown and
text have real lines, so they are cited as lines. A .docx has no lines at
all — what it has is a paragraph sequence — so those are cited as
paragraphs and the report says "paragraph", not "line". Calling a paragraph
a line would send someone to the wrong place in a Word document and make
them distrust everything else in the report.

**Some text is not a policy statement.** A term inside a fenced code block
or a YAML front-matter header is configuration or metadata, not a
commitment the organisation is making. Those regions are blanked rather
than deleted, so the surviving lines keep their original numbers.

**What could not be read has to be said out loud.** A folder with two
unreadable PDFs analysed as if it held only the readable files produces a
gap report that is wrong in the user's favour. Skips are collected and
reported, never swallowed.
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from .signatures import Document

READABLE = {".md", ".markdown", ".txt", ".rst", ".docx"}

# Anything above this is not a policy; it is an export or a log, and
# reading it wastes time and pollutes the result.
MAX_BYTES = 8 * 1024 * 1024

FENCE = re.compile(r"^\s*(```|~~~)")
FRONT_MATTER = re.compile(r"^---\s*$")


@dataclass(frozen=True)
class Skipped:
    """A file that was not read, and the reason a person can act on."""

    path: Path
    reason: str


@dataclass(frozen=True)
class LoadResult:
    documents: tuple[Document, ...]
    units: dict[str, str]
    skipped: tuple[Skipped, ...]

    @property
    def unit_for(self) -> dict[str, str]:
        return self.units


def _strip_markdown_noise(lines: list[str]) -> list[str]:
    """Blank out code fences and front matter, keeping line numbers intact."""
    cleaned = list(lines)
    in_fence = False
    in_front_matter = False

    for index, line in enumerate(lines):
        # Front matter only counts at the very top of the file.
        if index == 0 and FRONT_MATTER.match(line):
            in_front_matter = True
            cleaned[index] = ""
            continue
        if in_front_matter:
            cleaned[index] = ""
            if FRONT_MATTER.match(line):
                in_front_matter = False
            continue

        if FENCE.match(line):
            in_fence = not in_fence
            cleaned[index] = ""
            continue
        if in_fence:
            cleaned[index] = ""

    return cleaned


def _read_text(path: Path) -> list[str]:
    """Read a text file, tolerating the encodings real document sets carry."""
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1253", "iso-8859-7", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding).splitlines()
        except UnicodeDecodeError:
            continue
    # latin-1 cannot fail, so reaching here means the file is not text.
    raise ValueError("could not be decoded as text in any supported encoding")


def _read_docx(path: Path) -> list[str]:
    """Read a .docx as a paragraph sequence, including table cells.

    Tables matter more than they look. The document control block of a
    policy — owner, version, approval date, next review — is almost always
    a table, and dropping tables would blind every hygiene check that
    matters.
    """
    try:
        from docx import Document as DocxDocument
        from docx.opc.exceptions import OpcError
    except ImportError as error:  # pragma: no cover - environment dependent
        raise ValueError("reading .docx needs python-docx: pip install python-docx") from error

    # python-docx raises its own OpcError family for a file that is not a
    # valid package. A truncated or renamed file is common in a real
    # document set, and it must be reported rather than end the run.
    try:
        document = DocxDocument(str(path))
    except (OpcError, zipfile.BadZipFile, KeyError, ValueError) as error:
        raise ValueError(f"not a readable .docx ({type(error).__name__})") from error

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # A row reads as one line so that "Owner | Jane Doe" stays together.
            paragraphs.append(" | ".join(c for c in cells if c))
    return paragraphs


def load_document(path: Path, display_name: str | None = None) -> tuple[Document, str]:
    """Read one file. Returns the document and the unit its numbers refer to."""
    suffix = path.suffix.lower()
    name = display_name or path.name

    if suffix == ".docx":
        lines = _read_docx(path)
        unit = "paragraph"
    else:
        lines = _read_text(path)
        unit = "line"
        if suffix in {".md", ".markdown"}:
            lines = _strip_markdown_noise(lines)

    return Document(name=name, lines=tuple(lines)), unit


def load_directory(root: Path) -> LoadResult:
    """Read every readable document under a folder.

    Files are visited in sorted order so two runs over the same folder
    produce byte-identical output.
    """
    if not root.exists():
        raise FileNotFoundError(f"No such folder: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Not a folder: {root}")

    documents: list[Document] = []
    units: dict[str, str] = {}
    skipped: list[Skipped] = []

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue

        relative = path.relative_to(root)
        name = str(relative)

        if any(part.startswith(".") for part in relative.parts):
            continue
        # Word leaves these behind when a document is open.
        if path.name.startswith("~$"):
            continue

        suffix = path.suffix.lower()
        if suffix not in READABLE:
            skipped.append(
                Skipped(relative, f"{suffix or 'no extension'} is not a readable format")
            )
            continue

        size = path.stat().st_size
        if size == 0:
            skipped.append(Skipped(relative, "the file is empty"))
            continue
        if size > MAX_BYTES:
            skipped.append(
                Skipped(
                    relative,
                    f"{size // (1024 * 1024)} MB is larger than the {MAX_BYTES // (1024 * 1024)} MB limit",
                )
            )
            continue

        try:
            document, unit = load_document(path, display_name=name)
        except (ValueError, OSError) as error:
            skipped.append(Skipped(relative, str(error)))
            continue

        if not any(line.strip() for line in document.lines):
            skipped.append(Skipped(relative, "no readable text was found in it"))
            continue

        documents.append(document)
        units[name] = unit

    return LoadResult(tuple(documents), units, tuple(skipped))
